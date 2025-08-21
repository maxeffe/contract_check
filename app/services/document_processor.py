import io
import os
from typing import Optional, Tuple
from config.logging_config import prediction_logger

class DocumentProcessor:
    """Сервис для обработки различных типов документов"""
    
    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024
    
    def extract_text_from_pdf(self, file_content: bytes) -> Optional[str]:
        """Извлечение текста из PDF файла"""
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if text.strip():
                prediction_logger.info("Текст успешно извлечен из PDF")
                return text.strip()
            else:
                prediction_logger.warning("PDF файл не содержит извлекаемого текста")
                return None
                
        except ImportError:
            prediction_logger.error("PyPDF2 не установлен. Установите: pip install PyPDF2")
            return None
        except Exception as e:
            prediction_logger.error(f"Ошибка при обработке PDF: {str(e)}")
            return None
    
    def extract_text_from_docx(self, file_content: bytes) -> Optional[str]:
        """Извлечение текста из DOCX файла"""
        try:
            import docx
            
            doc_file = io.BytesIO(file_content)
            document = docx.Document(doc_file)
            
            text = ""
            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"
            
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                prediction_logger.info("Текст успешно извлечен из DOCX")
                return text.strip()
            else:
                prediction_logger.warning("DOCX файл не содержит текста")
                return None
                
        except ImportError:
            prediction_logger.error("python-docx не установлен. Установите: pip install python-docx")
            return None
        except Exception as e:
            prediction_logger.error(f"Ошибка при обработке DOCX: {str(e)}")
            return None
    
    def extract_text_from_doc(self, file_content: bytes) -> Optional[str]:
        """Извлечение текста из DOC файла (старый формат Word)"""
        try:
            import docx2txt
            
            doc_file = io.BytesIO(file_content)
            text = docx2txt.process(doc_file)
            
            if text and text.strip():
                prediction_logger.info("Текст успешно извлечен из DOC с помощью docx2txt")
                return text.strip()
            else:
                prediction_logger.warning("docx2txt не смог обработать DOC файл")
                
        except ImportError:
            prediction_logger.error("docx2txt не установлен. Установите: pip install docx2txt")
        except Exception as e:
            prediction_logger.warning(f"docx2txt не смог обработать DOC: {str(e)}")
        
        try:
            prediction_logger.info("Пробуем python-docx для DOC файла")
            text = self.extract_text_from_docx(file_content)
            if text and len(text.strip()) > 50:
                prediction_logger.info("DOC файл успешно обработан как DOCX")
                return text.strip()
        except Exception as e:
            prediction_logger.warning(f"python-docx не смог обработать DOC: {str(e)}")
        
        try:
            import zipfile
            
            doc_file = io.BytesIO(file_content)
            
            if zipfile.is_zipfile(doc_file):
                prediction_logger.info("DOC файл является ZIP-архивом, обрабатываем как DOCX")
                doc_file.seek(0)
                text = self.extract_text_from_docx(file_content)
                if text:
                    return text
            
            prediction_logger.info("Пробуем простое декодирование DOC файла")
            
            text_content = file_content.decode('latin-1', errors='ignore')
            
            import re
            
            prediction_logger.info("Ищем маркеры реального содержимого в DOC файле")
            
            contract_markers = [
                r'договор[а-я]*\s+[а-я]{3,}',  
                r'[а-я]{3,}\s+обязуется',     
                r'стороны\s+[а-я]{3,}',       
                r'исполнитель\s+[а-я]{3,}',    
                r'заказчик\s+[а-я]{3,}',       
                r'ответственность\s+[а-я]{3,}', 
                r'[а-я]{3,}\s+руб[лей]*',      
                r'пункт\s+\d+',                
                r'статья\s+\d+',               
                r'\d+\s+[а-я]{4,}\s+\d{4}\s*г', 
            ]
            

            contract_fragments = []
            for pattern in contract_markers:
                matches = re.findall(pattern, text_content, re.IGNORECASE)
                if matches:
                    prediction_logger.info(f"Найден маркер договора: {pattern}")
                    for match in matches:
                        match_pos = text_content.lower().find(match.lower())
                        if match_pos >= 0:
                            start = max(0, match_pos - 200)
                            end = min(len(text_content), match_pos + 200)
                            context = text_content[start:end]
                            contract_fragments.append(context)
            
            if contract_fragments:
                combined_text = ' '.join(contract_fragments)
                prediction_logger.info(f"Найдено {len(contract_fragments)} фрагментов договора")
            else:
                combined_text = text_content
                prediction_logger.info("Маркеры договора не найдены, используем полный текст")
            
            text_without_xml = re.sub(r'<[^>]*>', ' ', combined_text)
            
            office_junk_patterns = [
                r'xmlns[:\w]*[=\s]*["\'][^"\']*["\']',
                r'schemas\.[A-Za-z0-9./-]+',
                r'openxmlformats\.[A-Za-z0-9./-]*',
                r'microsoft\.com[A-Za-z0-9./-]*',
                r'w3\.org[A-Za-z0-9./-]*',
                r'DocumentLibraryForm[A-Za-z0-9\s]*',
                r'themeManager\.xml[A-Za-z0-9\s]*',
                r'rels\.rels[A-Za-z0-9\s]*',
                r'PK\s*-\s*!\s*[A-Za-z0-9\s]*',
                r'revisions?\.\s*[A-Za-z\s]*',
                r'saves?\s+or\s+revisions?',
                r'application\s+is\s+responsible',
                r'This\s+value\s+indicates',
                r'number\s+of\s+saves',
            ]
            
            cleaned_text = text_without_xml
            for pattern in office_junk_patterns:
                cleaned_text = re.sub(pattern, ' ', cleaned_text, flags=re.IGNORECASE)
            
            office_words = [
                'Document', 'Summary', 'Information', 'officeDocument', 
                'clrMap', 'drawingml', 'DocumentLibraryForm', 'themeManager',
                'rels', 'theme', 'responsible', 'updating', 'revision',
                'indicates', 'value', 'saves', 'application'
            ]
            
            for word in office_words:
                cleaned_text = re.sub(rf'\b{re.escape(word)}\b', ' ', cleaned_text, flags=re.IGNORECASE)
            
            cleaned_text = re.sub(r'[!@#$%^&*()_+=\[\]{}|\\:";\'<>?,./`~]{3,}', ' ', cleaned_text)
            cleaned_text = re.sub(r'[\x00-\x1f\x7f-\x9f]', ' ', cleaned_text)
            
            cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
            cleaned_text = re.sub(r'\s*([.,:;!?])\s*', r'\1 ', cleaned_text)
            
            meaningful_parts = []
            
            russian_parts = re.findall(r'[А-Яа-я][а-я\s\.,!?;:()\-№%]{20,}', cleaned_text)
            meaningful_parts.extend(russian_parts)
            
            english_parts = re.findall(r'[A-Z][A-Za-z\s\.,!?;:()\-]{20,}', cleaned_text)  
            meaningful_parts.extend(english_parts)
            
            numeric_parts = re.findall(r'[а-яА-Я\s]{5,}[\d\s,%]{2,}[а-яА-Я\s]{5,}', cleaned_text)
            meaningful_parts.extend(numeric_parts)
            
            if meaningful_parts:
                unique_parts = list(set(meaningful_parts))
                unique_parts = [part.strip() for part in unique_parts if len(part.strip()) > 15]
                unique_parts.sort(key=len, reverse=True)
                
                result_text = ' '.join(unique_parts[:5])
                result_text = re.sub(r'\s+', ' ', result_text).strip()

                if (len(result_text) > 100 and 
                    not re.search(r'(PK|xml|rels|theme|Form)', result_text, re.IGNORECASE) and
                    re.search(r'[а-яА-Я]{10,}', result_text)):
                    
                    prediction_logger.info(f"Качественный текст извлечен из DOC: {len(result_text)} символов")
                    return result_text
                else:
                    prediction_logger.warning(f"Извлеченный текст низкого качества: {result_text[:100]}...")
            
            prediction_logger.warning("Не удалось извлечь содержательный текст из DOC файла")
            
            return "Документ содержит преимущественно служебную информацию. Невозможно извлечь содержательный текст для анализа. Рекомендуется использовать файл в формате DOCX или TXT."
            
            prediction_logger.warning("Не удалось извлечь содержательный текст из DOC файла")
            return None
            
        except Exception as e:
            prediction_logger.error(f"Критическая ошибка при обработке DOC: {str(e)}")
            return None
    
    def extract_text_from_txt(self, file_content: bytes) -> Optional[str]:
        """Извлечение текста из TXT файла"""
        try:
            encodings = ['utf-8', 'cp1251', 'latin-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    if text.strip():
                        prediction_logger.info(f"Текст успешно извлечен из TXT (кодировка: {encoding})")
                        return text.strip()
                except UnicodeDecodeError:
                    continue
            
            prediction_logger.error("Не удалось декодировать TXT файл ни в одной из поддерживаемых кодировок")
            return None
            
        except Exception as e:
            prediction_logger.error(f"Ошибка при обработке TXT: {str(e)}")
            return None
    
    def process_file(self, file_content: bytes, filename: str) -> Tuple[Optional[str], bool]:
        """
        Обработать файл и извлечь текст
        
        Returns:
            Tuple[Optional[str], bool]: (извлеченный_текст, успешно_обработан)
        """
        if len(file_content) > self.max_file_size:
            prediction_logger.error(f"Файл {filename} слишком большой ({len(file_content)} байт)")
            return None, False
        
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            text = self.extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            text = self.extract_text_from_docx(file_content)
        elif filename_lower.endswith('.doc'):
            text = self.extract_text_from_doc(file_content)
        elif filename_lower.endswith('.txt'):
            text = self.extract_text_from_txt(file_content)
        else:
            prediction_logger.error(f"Неподдерживаемый тип файла: {filename}")
            return None, False
        
        if text and len(text.strip()) > 10:
            return text, True
        else:
            prediction_logger.warning(f"Из файла {filename} не удалось извлечь достаточно текста")
            return None, False
    
    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, str]:
        """
        Валидация файла
        
        Returns:
            Tuple[bool, str]: (валиден, сообщение_об_ошибке)
        """
        if not file_content:
            return False, "Файл пустой"
        
        if len(file_content) > self.max_file_size:
            return False, f"Файл слишком большой (максимум {self.max_file_size // 1024 // 1024} MB)"
        
        supported_extensions = ['.txt', '.pdf', '.doc', '.docx']
        file_extension = None
        
        for ext in supported_extensions:
            if filename.lower().endswith(ext):
                file_extension = ext
                break
        
        if not file_extension:
            return False, f"Неподдерживаемый тип файла. Поддерживаемые: {', '.join(supported_extensions)}"
        
        return True, "OK"
    
    def get_required_packages(self) -> dict:
        """Получить список требуемых пакетов для обработки файлов"""
        return {
            'pdf': 'PyPDF2',
            'docx': 'python-docx', 
            'doc': 'docx2txt'
        }

document_processor = DocumentProcessor()